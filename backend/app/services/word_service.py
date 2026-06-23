"""Build an SMK Word (.docx) document with python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# (번호, 제목) 순서
_SECTIONS = [
    ("01", "기술개요", "item_01"),
    ("02", "기술의 차별성", "item_02"),
    ("03", "주요 키워드", "item_03"),
    ("04", "TRL 단계", "item_04"),
    ("05", "사업화 포인트", "item_05"),
    ("06", "활용분야", "item_06"),
]

_HEADER_ROW = [
    "구분",
    "특허명",
    "등록/공개",
    "공보/등록번호",
    "출원번호",
    "출원인",
    "출원일",
]


def _safe(value: str, default: str) -> str:
    value = (value or "").strip()
    if not value:
        return default
    # 파일명에 쓸 수 없는 문자 제거.
    return re.sub(r'[\\/:*?"<>|]+', "_", value)


def _add_heading(doc: Document, text: str, size: int = 13) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x8A, 0x1C, 0x3B)


def _add_body(doc: Document, body: str) -> None:
    for line in (body or "").split("\n"):
        line = line.strip()
        if not line:
            continue
        doc.add_paragraph(line)


def build_smk_docx(
    patent: dict,
    items: dict,
    ip_table: list[dict],
    output_dir: Path,
    market: dict | None = None,
    sources: list[str] | None = None,
) -> tuple[Path, str]:
    """Create the SMK .docx and return (path, filename)."""
    doc = Document()

    # 제목 / 부제
    title = doc.add_heading("기술이전 마케팅 자료 (SMK)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(patent.get("title") or "(발명의 명칭 없음)")
    sub_run.bold = True
    sub_run.font.size = Pt(14)

    # 구분선
    divider = doc.add_paragraph("─" * 46)
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 01 ~ 06
    for no, sec_title, key in _SECTIONS:
        _add_heading(doc, f"{no}. {sec_title}")
        _add_body(doc, items.get(key, ""))

    # 07 시장규모
    _add_heading(doc, "07. 시장규모")
    if market:
        summary = market.get("summary") or market.get("text") or ""
        if summary:
            _add_body(doc, summary)
        if market.get("has_chart") and market.get("years"):
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "연도"
            hdr[1].text = "국내 (억원)"
            hdr[2].text = "세계 (억원)"
            years = market.get("years") or []
            domestic = market.get("domestic") or []
            global_vals = market.get("global_values") or market.get("global") or []
            for i, year in enumerate(years):
                cells = table.add_row().cells
                cells[0].text = str(year)
                cells[1].text = str(domestic[i]) if i < len(domestic) else ""
                cells[2].text = str(global_vals[i]) if i < len(global_vals) else ""
        elif market.get("rows"):
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "연도"
            hdr[1].text = "시장규모"
            for row in market["rows"]:
                cells = table.add_row().cells
                cells[0].text = str(row.get("year", ""))
                cells[1].text = str(row.get("value", ""))
    else:
        _add_body(doc, "시장규모 데이터가 아직 준비되지 않았습니다.")

    # 08 지식재산권 현황 (표)
    _add_heading(doc, "08. 지식재산권 현황")
    table = doc.add_table(rows=1, cols=len(_HEADER_ROW))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _HEADER_ROW):
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in ip_table or []:
        cells = table.add_row().cells
        cells[0].text = row.get("category", "") or ""
        cells[1].text = row.get("title", "") or ""
        cells[2].text = row.get("doc_type", "") or ""
        cells[3].text = row.get("doc_no", "") or ""
        cells[4].text = row.get("app_no", "") or ""
        cells[5].text = row.get("applicant", "") or ""
        cells[6].text = row.get("apply_date", "") or ""

    # 출처 목록
    if sources:
        _add_heading(doc, "출처")
        for src in sources:
            doc.add_paragraph(str(src), style="List Bullet")

    output_dir.mkdir(parents=True, exist_ok=True)
    filename = (
        f"SMK_{_safe(patent.get('app_no'), 'unknown')}"
        f"_{_safe(patent.get('apply_date'), 'nodate')}.docx"
    )
    path = output_dir / filename
    doc.save(path)
    return path, filename
